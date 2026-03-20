#!/usr/bin/env python3
"""
Extract content from DOCX whitepaper files for HTML conversion
"""
import sys
import os

try:
    from docx import Document
except ImportError:
    print("python-docx not installed. Installing...")
    os.system("pip install python-docx")
    from docx import Document


def extract_docx_content(docx_path, output_txt_path):
    """Extract text content from DOCX file and save to text file"""
    try:
        doc = Document(docx_path)

        with open(output_txt_path, 'w', encoding='utf-8') as f:
            f.write(f"=== EXTRACTED FROM: {os.path.basename(docx_path)} ===\n\n")

            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:
                    # Identify heading levels based on style
                    style_name = para.style.name.lower() if para.style else ''

                    if 'heading 1' in style_name or 'title' in style_name:
                        f.write(f"\n# {text}\n\n")
                    elif 'heading 2' in style_name:
                        f.write(f"\n## {text}\n\n")
                    elif 'heading 3' in style_name:
                        f.write(f"\n### {text}\n\n")
                    else:
                        f.write(f"{text}\n\n")

            # Extract tables
            if doc.tables:
                f.write("\n\n=== TABLES ===\n\n")
                for table_idx, table in enumerate(doc.tables):
                    f.write(f"\nTable {table_idx + 1}:\n")
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells)
                        f.write(f"{row_text}\n")
                    f.write("\n")

        print(f"✓ Extracted: {docx_path} -> {output_txt_path}")
        return True

    except Exception as e:
        print(f"✗ Error extracting {docx_path}: {e}")
        return False


def main():
    # Files to extract
    files = [
        "Speeki_Whitepaper_ICSR_Internal_Controls_Sustainability_Reporting.docx",
        "Employee_Hiring_in_ISO_37001.docx",
        "ABMS_Risk_Assessment_Whitepaper.docx",
        "speeki_ncr_whitepaper.docx",
        "Beyond_the_Checkbox_ABMS_Objectives_Article.docx"
    ]

    base_path = "/tmp/checklists"
    output_dir = os.path.join(base_path, "extracted_content")

    # Create output directory
    os.makedirs(output_dir, exist_ok=True)

    print("\n📄 Extracting DOCX content...\n")

    for filename in files:
        docx_path = os.path.join(base_path, filename)
        output_name = filename.replace('.docx', '.txt')
        output_path = os.path.join(output_dir, output_name)

        if os.path.exists(docx_path):
            extract_docx_content(docx_path, output_path)
        else:
            print(f"⚠ File not found: {docx_path}")

    print(f"\n✓ All content extracted to: {output_dir}/\n")


if __name__ == "__main__":
    main()
